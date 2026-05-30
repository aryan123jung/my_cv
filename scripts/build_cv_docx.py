from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "exports"
OUT_DIR.mkdir(exist_ok=True)
DOCX_OUT = OUT_DIR / "Aryan_Jung_Chhetri_CV.docx"
PHOTO = ROOT / "public" / "IMG_1907.jpg"

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(18, 31, 46)
MUTED = RGBColor(88, 98, 110)
LIGHT_FILL = "E8EEF5"
PALE_FILL = "F4F7FB"
WHITE = "FFFFFF"
BORDER = "D7DEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    existing_grid = table._tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.2, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph(p, before=0, after=4, line=1.18, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(p, text, size=10.2, color=DARK, bold=False, italic=False):
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_hyperlink(paragraph, text, url, size=9.6):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    for r in paragraph.runs:
        if r.text == text:
            set_run_font(r, size=size, color=ACCENT)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=7, after=3, line=1.0)
    r = add_text(p, text.upper(), size=9.5, color=ACCENT, bold=True)
    r.font.all_caps = True
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D7DEE8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc, text, bold_prefix=None, after=2):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=after, line=1.12)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=9.6, color=DARK, bold=True)
        add_text(p, text[len(bold_prefix):], size=9.6, color=DARK)
    else:
        add_text(p, text, size=9.6, color=DARK)
    return p


def add_chip_table(doc, items, columns=3):
    rows = (len(items) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    widths = [3120, 3120, 3120][:columns]
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, PALE_FILL)
            set_cell_border(cell, "E2E8F0", "4")
            set_cell_margins(cell, top=70, bottom=70, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for idx, item in enumerate(items):
        cell = table.rows[idx // columns].cells[idx % columns]
        p = cell.paragraphs[0]
        set_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, item, size=9.2, color=DARK, bold=True)
    for idx in range(len(items), rows * columns):
        cell = table.rows[idx // columns].cells[idx % columns]
        set_cell_shading(cell, WHITE)
        p = cell.paragraphs[0]
        set_paragraph(p, before=0, after=0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.2)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.18

for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.color.rgb = ACCENT
    style.font.bold = True

footer = section.footer.paragraphs[0]
set_paragraph(footer, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(footer, "Aryan Jung Chhetri | Curriculum Vitae", size=8.5, color=MUTED)

header_table = doc.add_table(rows=1, cols=2)
set_table_width(header_table, [7800, 1560])
for cell in header_table.rows[0].cells:
    set_cell_shading(cell, WHITE)
    set_cell_border(cell, WHITE, "0")
    set_cell_margins(cell, top=0, bottom=90, start=0, end=0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

left, right = header_table.rows[0].cells
p = left.paragraphs[0]
set_paragraph(p, before=0, after=2, line=1.0)
add_text(p, "ARYAN JUNG CHHETRI", size=21.5, color=DARK, bold=True)

p = left.add_paragraph()
set_paragraph(p, before=0, after=5, line=1.05)
add_text(
    p,
    "Fullstack Developer | Mobile App Developer | UI/UX Enthusiast | Computing Student",
    size=10.5,
    color=ACCENT,
    bold=True,
)

p = left.add_paragraph()
set_paragraph(p, before=0, after=0, line=1.22)
add_text(p, "Dhobighat, Lalitpur | ", size=9.5, color=MUTED)
add_text(p, "chhetriaryanjung@gmail.com", size=9.5, color=MUTED)
add_text(p, " | 9840530090", size=9.5, color=MUTED)

p = left.add_paragraph()
set_paragraph(p, before=1, after=0, line=1.18)
add_hyperlink(p, "GitHub", "https://github.com/aryan123jung", 9.3)
add_text(p, " | ", size=9.3, color=MUTED)
add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/aryan-jung-chhetri-b67049363/", 9.3)

if PHOTO.exists():
    p = right.paragraphs[0]
    set_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run().add_picture(str(PHOTO), width=Inches(1.0))

add_section_heading(doc, "Professional Profile")
p = doc.add_paragraph()
set_paragraph(p, before=0, after=4, line=1.18)
add_text(
    p,
    "BSc (Hons) Computing student with a practical blend of fullstack development, "
    "mobile app development, UI/UX design, cybersecurity learning, and data science "
    "and analytics. Builds clean, accessible, recruiter-ready digital experiences "
    "with strong attention to visual hierarchy, usability, and secure web fundamentals.",
    size=10.1,
)

add_section_heading(doc, "Technical Skills")
add_chip_table(
    doc,
    [
        "React",
        "JavaScript",
        "Flutter",
        "Figma",
        "Burp Suite",
        "Python",
        "Java",
        "HTML",
        "CSS",
        "UI/UX Design",
        "Data Analytics",
        "Web Security",
    ],
    columns=3,
)

add_section_heading(doc, "Education")
p = doc.add_paragraph()
set_paragraph(p, before=0, after=1, line=1.12)
add_text(p, "Softwarica College of IT and E-Commerce: ", size=10.0, color=DARK, bold=True)
add_text(p, "BSc (Hons) Computing", size=9.7, color=DARK)
p = doc.add_paragraph()
set_paragraph(p, before=0, after=0, line=1.12)
add_text(p, "+2 Nightingale College: ", size=10.0, color=DARK, bold=True)
add_text(p, "Management with Maths and Computer", size=9.7, color=DARK)

add_section_heading(doc, "Projects")
projects = [
    (
        "ChautariKuraKani",
        "Full social media application for web and mobile where users connect and join Chautari communities inspired by Nepali conversation culture. Tech: TypeScript, Flutter, NodeJS.",
    ),
    (
        "TeamSphere",
        "Team management platform for coaches to analyze player performance and support structured decisions. Tech: React, Fullstack, Tailwind.",
    ),
    (
        "Caption",
        "Mobile application for multi-language translation and scan-based translation workflows. Tech: Kotlin, Java.",
    ),
    (
        "GuffGuthi",
        "Reddit-like communication app inspired by Nepali Guthi culture, focused on community discussion and user interaction. Tech: React, JavaScript.",
    ),
    (
        "Data Science Data Analytics",
        "Analytics project for a UK county dataset, focused on insight generation and informed decision-making. Tech: R, R-Studio, Excel.",
    ),
    (
        "ML/AI Player Recommendation System",
        "ML/AI project for IPL franchises to support squad building and identify domestic players with growth potential. Tech: Python, Excel.",
    ),
]
for name, detail in projects:
    add_bullet(doc, f"{name}: {detail}", bold_prefix=f"{name}:")

add_section_heading(doc, "Academic & Personal Experience")
experience = [
    ("Personal Web Application Development", "Built fullstack interfaces and practiced component-driven development with modern JavaScript workflows."),
    ("Personal Mobile Application Development", "Built functional mobile interfaces using modern Flutter and Riverpod workflows."),
    ("UI/UX Design Practice", "Created clean user flows, wireframes, and interface concepts with focus on usability and visual hierarchy."),
    ("Academic Team Collaborations", "Worked with classmates on computing assignments, shared responsibilities, and presented technical outcomes."),
    ("Cybersecurity Learning", "Explored web application security fundamentals, HTTP behavior, and testing workflows using Burp Suite."),
]
for name, detail in experience:
    add_bullet(doc, f"{name}: {detail}", bold_prefix=f"{name}:")

add_section_heading(doc, "Leadership & Soft Skills")
add_bullet(
    doc,
    "U19 District Cricket Captain: Demonstrated leadership, teamwork, discipline, communication, and decision-making under pressure.",
    bold_prefix="U19 District Cricket Captain:",
)
p = doc.add_paragraph()
set_paragraph(p, before=0, after=0, line=1.12)
add_text(p, "Soft skills: ", size=9.6, color=DARK, bold=True)
add_text(
    p,
    "Leadership, Team Collaboration, Communication, Problem Solving, Adaptability, Strategic Thinking",
    size=9.6,
    color=DARK,
)

doc.core_properties.author = "Aryan Jung Chhetri"
doc.core_properties.title = "Aryan Jung Chhetri CV"
doc.core_properties.subject = "Curriculum Vitae"
doc.core_properties.keywords = "Fullstack Developer, Mobile App Developer, UI/UX, Cybersecurity, Data Analytics"
doc.save(DOCX_OUT)
print(DOCX_OUT)
